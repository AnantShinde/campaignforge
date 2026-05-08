"""
generate_doc.py — produces CampaignForge_Design.docx with embedded images.
Run from project root: python scripts/generate_doc.py
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
RESULTS = os.path.join(ROOT, "results")
OUT = os.path.join(ROOT, "CampaignForge_Design.docx")

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Helpers ───────────────────────────────────────────────────────────────────
def h1(text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    return p

def h2(text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = RGBColor(0x16, 0x21, 0x3e)
    return p

def h3(text):
    return doc.add_heading(text, level=3)

def body(text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(11)
    return p

def code(text):
    p = doc.add_paragraph()
    p.style = doc.styles['No Spacing']
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1f, 0x1f, 0x1f)
    p.paragraph_format.left_indent = Cm(0.8)
    return p

def add_image(path, width_inches=6.0, caption=None):
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(width_inches))
        last = doc.paragraphs[-1]
        last.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if caption:
            c = doc.add_paragraph(caption)
            c.alignment = WD_ALIGN_PARAGRAPH.CENTER
            c.runs[0].font.size = Pt(9)
            c.runs[0].italic = True
    else:
        body(f"[Image not found: {path}]", italic=True)

def add_table(headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Light List Accent 1"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for run in hdr[i].paragraphs[0].runs:
            run.bold = True
    for ri, row in enumerate(rows):
        cells = t.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = val
    doc.add_paragraph()

def divider():
    doc.add_paragraph("─" * 80)

# ═══════════════════════════════════════════════════════════════════════════════
# COVER
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_heading("CampaignForge", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub = doc.add_paragraph("System Design & Implementation")
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.runs[0].font.size = Pt(16)
sub.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55)
doc.add_paragraph()
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 1. PROBLEM STATEMENT
# ═══════════════════════════════════════════════════════════════════════════════
h1("1. Problem Statement")
body(
    "A global consumer goods company launches hundreds of localised social ad campaigns "
    "every month across dozens of markets. Today this is done manually:"
)
for item in [
    "A creative team writes briefs",
    "Agencies produce assets per region",
    "Legal and brand teams review",
    "Stakeholders in each market approve",
    "Assets get scheduled and published",
]:
    p = doc.add_paragraph(item, style="List Bullet")
    p.runs[0].font.size = Pt(11)

body(
    "This process is slow (weeks per campaign), expensive (agencies + revisions), "
    "inconsistent (off-brand creative slips through), and impossible to learn from "
    "(siloed performance data)."
)
body(
    "My goal: design and build a system that takes a campaign brief as input and produces "
    "brand-compliant, legally-vetted, localised creative assets for three social platforms "
    "— fully automated, under 30 seconds.",
    bold=True
)

# ═══════════════════════════════════════════════════════════════════════════════
# 2. REQUIREMENTS
# ═══════════════════════════════════════════════════════════════════════════════
h1("2. Requirements")
h2("Functional Requirements")
add_table(
    ["ID", "Requirement"],
    [
        ["FR1", "Accept a campaign brief (JSON) with product name, region, audience, message, language"],
        ["FR2", "Generate images in 1:1 (Instagram), 9:16 (TikTok/Reels), 16:9 (YouTube)"],
        ["FR3", "Overlay localised text on each generated image"],
        ["FR4", "Save assets at outputs/{campaign_id}/{product_name}/{ratio}.png"],
        ["FR5", "Enforce brand compliance on all generated creative"],
        ["FR6", "Enforce legal compliance — no prohibited claims or terms"],
        ["FR7", "Approval workflow — campaigns land in pending_review before going final"],
        ["FR8", "Batch submission — up to 50 briefs per request"],
        ["FR9", "Programmatic download via S3 CORS and Blob fetching"],
    ]
)
h2("Non-Functional Requirements")
add_table(
    ["ID", "Requirement"],
    [
        ["NFR1", "Scale horizontally without infrastructure changes"],
        ["NFR2", "Generation of one campaign must complete within 10 minutes"],
        ["NFR3", "Brand and legal guardrails must apply to 100% of outputs"],
        ["NFR4", "Reproducible across environments via Infrastructure as Code"],
        ["NFR5", "All generated assets must be durable — S3 with versioning enabled"],
        ["NFR6", "Generation must be idempotent — SQS retries must not duplicate work"],
    ]
)

# ═══════════════════════════════════════════════════════════════════════════════
# 3. ROUGH ESTIMATES
# ═══════════════════════════════════════════════════════════════════════════════
h1("3. Rough Estimates")
h2("Volume")
for line in [
    "Campaigns per month: ~500",
    "Images per campaign: 3 aspect ratios",
    "Total images per month: 500 × 3 = 1,500 images",
]:
    doc.add_paragraph(line, style="List Bullet").runs[0].font.size = Pt(11)

h2("Compute & Throughput")
for line in [
    "Imagen 4 Fast generation: ~8–12s per image",
    "3 images fired concurrently → ~12s wall-clock for images",
    "GPT-4o mini (copy + 2 compliance passes): ~8s, runs in parallel",
    "Total wall-clock: ~20–30s per campaign",
]:
    doc.add_paragraph(line, style="List Bullet").runs[0].font.size = Pt(11)

h2("Cost per Campaign")
add_table(
    ["Component", "Per Campaign", "Per Month (500)"],
    [
        ["Google Imagen 4 Fast (3 images)", "~$0.12", "~$60"],
        ["GPT-4o mini (copy + compliance ×2)", "~$0.006", "~$3"],
        ["S3 + DynamoDB + Lambda", "~$0.001", "negligible"],
        ["Total", "~$0.13", "~$65"],
    ]
)

# ═══════════════════════════════════════════════════════════════════════════════
# 4. HIGH-LEVEL DESIGN
# ═══════════════════════════════════════════════════════════════════════════════
h1("4. High-Level Design")
body("I organised the system into five layers, each with a single responsibility:")
add_table(
    ["Layer", "Components", "Responsibility"],
    [
        ["P1 User", "API Gateway HTTP v2 + Cognito", "Authenticated brief ingestion"],
        ["P2 Queue", "SQS Standard", "Decoupling + idempotent retry (NFR6)"],
        ["P3 Generation", "GenerateCampaign Lambda", "Copy, images, compliance, persistence"],
        ["P4 AI", "GPT-4o mini + Google Imagen 4", "Language model + image model"],
        ["P5 Storage", "S3 + DynamoDB", "Durable asset storage + campaign state"],
    ]
)

body("Pipeline flow:")
for step in [
    "POST /brief → 202 Accepted + campaign_id",
    "SQS receives the message (NFR6: idempotent, maxConcurrency=10)",
    "GenerateCampaign Lambda fires: CopyGen ‖ Pass1 Compliance (parallel)",
    "Pass 1 fail? → compliance_blocked, stop. Saves ~$0.12 in image costs.",
    "Imagen 4 Fast generates all 3 ratios concurrently",
    "Pillow composites the localised headline onto each image",
    "Pass 2 vision compliance checks the rendered output",
    "DynamoDB + S3 write → approval_status: pending_review",
    "GET /campaigns/{id} polls until complete, returns assets + report",
]:
    p = doc.add_paragraph(step, style="List Number")
    p.runs[0].font.size = Pt(11)

# ═══════════════════════════════════════════════════════════════════════════════
# 5. HOW IT WORKS
# ═══════════════════════════════════════════════════════════════════════════════
h1("5. How It Works")

h2("5.A Two-Pass Compliance")
body(
    "The most important architectural decision I made was splitting compliance into two passes "
    "with a deliberate abort gate between them."
)
h3("Pass 1 — Pre-generation text check (~3s, parallel with CopyGen)")
body(
    "I check the brief's core message before any image is generated. "
    "If overall = fail, the pipeline stops — no Imagen 4 calls. "
    "This saves ~$0.12 per rejected campaign. At 500 campaigns/month, "
    "even a 5% fail rate saves $3/month and keeps the pipeline clean."
)
add_table(
    ["Criterion", "What I check"],
    [
        ["prohibited_claims", "No 'guaranteed', '#1', 'proven', 'cure', 'risk-free'"],
        ["legal_disclaimer", "Claims modest enough or disclaimer present"],
        ["brand_voice", "Premium, aspirational, authentic tone"],
        ["cultural_sensitivity", "Appropriate for the target region"],
        ["pii_risk", "No personal data, phone numbers, URLs"],
    ]
)
h3("Pass 2 — Post-generation vision check (~5s, after text overlay)")
body(
    "GPT-4o mini reviews the rendered composite image — not the text prompt, "
    "the actual generated PNG. This catches visual brand violations that text analysis "
    "cannot see: off-colour renders, illegible overlays, inappropriate generated content."
)
add_table(
    ["Criterion", "What I check"],
    [
        ["brand_colors", "Dominant palette is professional and consistent"],
        ["text_legibility", "Overlaid headline is clearly readable"],
        ["product_prominence", "Product is visible and well-presented"],
        ["visual_appropriateness", "Image is suitable for all audiences"],
    ]
)
body(
    "A Pass 2 fail sets approval_status = compliance_blocked. "
    "Human override requires the compliance_override Cognito group."
)

h2("5.B Component Design")
body("The entire generation pipeline runs inside a single Lambda — one SQS message, one execution.")
code(
    "generate_campaign/\n"
    "  main.py          — SQS handler + orchestrator\n"
    "  copy_gen.py      — GPT-4o mini: ad copy + image prompt\n"
    "  compliance.py    — GPT-4o mini: Pass 1 (text) + Pass 2 (vision)\n"
    "  image_gen.py     — Google Imagen 4 Fast, 3 ratios concurrent\n"
    "  text_overlay.py  — Pillow: headline compositing\n"
    "  config.py        — boto3 clients + cached API clients"
)
body(
    "Internal concurrency uses ThreadPoolExecutor at two levels: "
    "CopyGen ‖ Pass1 run in parallel (~8s wall-clock instead of ~16s), "
    "and all 3 image ratios generate concurrently (~12s instead of ~36s). "
    "Total wall-clock saving: ~60% vs sequential execution."
)

h2("5.C Database Design")
body(
    "I chose DynamoDB with a composite primary key: campaign_id (PK) + product_name (SK). "
    "This means one Query retrieves all products for a campaign, and one GetItem targets "
    "a specific product. No scan needed for either access pattern."
)
add_table(
    ["Attribute", "Type", "Notes"],
    [
        ["campaign_id", "String", "UUID — partition key"],
        ["product_name", "String", "Sort key"],
        ["region", "String", "e.g. 'brazil'"],
        ["language", "String", "BCP-47 code, e.g. 'pt-BR'"],
        ["ad_copy", "List", "[{ lang, headline, body, cta }]"],
        ["images", "Map", '{ "1x1": { url, s3_key }, "9x16": {...}, "16x9": {...} }'],
        ["compliance_pass1", "Map", "GPT-4o mini text check result"],
        ["compliance_pass2", "Map", "GPT-4o mini vision check result"],
        ["approval_status", "String", "pending_review | approved | rejected | compliance_blocked"],
        ["ttl", "Number", "Unix timestamp — auto-expire after 90 days"],
    ]
)
body(
    "The GSI (status-created-index) powers the reviewer dashboard: "
    "query all campaigns WHERE approval_status = 'pending_review' "
    "ORDER BY created_at DESC — without a full table scan. "
    "The GSI is sparse: items still generating (no approval_status) are excluded automatically."
)

h2("5.D API Design")
body("Four endpoints, all protected by Cognito JWT (Authorization: Bearer <id_token>).")
add_table(
    ["Method", "Path", "Description"],
    [
        ["POST", "/brief", "Submit a single campaign brief → 202 + campaign_id"],
        ["POST", "/brief/batch", "Submit up to 50 briefs → 202 + campaign_ids[]"],
        ["GET", "/campaigns", "List all campaigns (history view)"],
        ["GET", "/campaigns/{id}", "Poll for status + presigned asset URLs"],
    ]
)

h2("5.E AI Pipeline")
add_table(
    ["Model", "Role", "Provider"],
    [
        ["gpt-4o-mini", "Ad copy generation, localisation, 2-pass compliance", "OpenAI"],
        ["imagen-4.0-fast-generate-001", "Image generation — exact 1:1, 9:16, 16:9", "Google AI"],
    ]
)
body(
    "I evaluated every accessible image generation model before settling on Imagen 4 Fast. "
    "Nova Canvas requires Bedrock model access activation. DALL-E 3 has no native 9:16. "
    "Stability models require an AWS Marketplace subscription. "
    "Imagen 4 Fast supports exact aspect ratios, costs ~$0.04/image, and was immediately accessible."
)

# ═══════════════════════════════════════════════════════════════════════════════
# 6. RESULTS
# ═══════════════════════════════════════════════════════════════════════════════
h1("6. Results")
h2("Ergo Desk Pro — United States")
body("Generated in 19 seconds. Compliance: pass / pass. Approval: pending_review.")

ergo_dir = os.path.join(RESULTS, "test_20260505_230810", "images")
add_image(os.path.join(ergo_dir, "1-1.png"),  width_inches=2.8, caption="Instagram 1×1")
add_image(os.path.join(ergo_dir, "9-16.png"), width_inches=1.8, caption="TikTok/Reels 9×16")
add_image(os.path.join(ergo_dir, "16-9.png"), width_inches=4.5, caption="YouTube 16×9")

doc.add_paragraph()
h2("LuxeRoll Hair Brush — Multi-Market Batch")
body(
    "I submitted 5 briefs simultaneously targeting working women aged 24–40 across "
    "4 global markets. All generated concurrently. Total time under 35 seconds."
)

hair_dir = os.path.join(RESULTS, "hairbrush_1778039117")
markets = [
    ("united_states", "United States — African American women"),
    ("brazil",        "Brazil — Latina women (copy in Portuguese)"),
    ("india",         "India — South Asian professional women"),
    ("germany",       "Germany — European professional women (copy in German)"),
]
for folder, label in markets:
    h3(label)
    mdir = os.path.join(hair_dir, folder)
    for ratio, caption, width in [
        ("1x1.png",  "Instagram 1×1", 2.4),
        ("9x16.png", "TikTok 9×16",   1.6),
        ("16x9.png", "YouTube 16×9",  3.8),
    ]:
        add_image(os.path.join(mdir, ratio), width_inches=width, caption=caption)
    doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
# 7. WHY I BUILT IT THIS WAY
# ═══════════════════════════════════════════════════════════════════════════════
h1("7. Why I Built It This Way")

decisions = [
    (
        "Single Lambda vs Step Functions",
        "Single GenerateCampaign Lambda with internal ThreadPoolExecutor",
        "AWS Step Functions Express Workflow (7 states)",
        "Step Functions adds 7 IAM roles, an EventBridge Pipe, and per-state CloudWatch streams "
        "with no material benefit at POC scale. The single Lambda achieves the same parallelism "
        "in ~20s. Step Functions is the right call if individual states need independent retry "
        "budgets or if the pipeline exceeds Lambda's 15-minute limit.",
    ),
    (
        "Two-Pass Compliance vs Single-Pass Parallel",
        "Sequential two passes: text pre-generation + vision post-generation",
        "Single compliance pass running parallel with image generation",
        "Pass 1 running before image generation means a hard failure aborts before any "
        "Imagen 4 calls — saving ~$0.12 per rejected campaign. Pass 2 vision-checks the "
        "rendered composite, catching visual violations invisible from text alone. "
        "Net latency cost is near zero because Pass 1 runs in parallel with CopyGen.",
    ),
    (
        "Google Imagen 4 vs Alternatives",
        "imagen-4.0-fast-generate-001 via Google AI API",
        "DALL-E 3 (OpenAI), Nova Canvas (Bedrock), Stability (Bedrock)",
        "I tested every accessible image generation model. "
        "Nova Canvas and Titan require Bedrock model access activation. "
        "Stability requires an AWS Marketplace subscription. "
        "DALL-E 3 lacks native 9:16. Imagen 4 Fast supports exact 1:1, 9:16, 16:9 natively, "
        "costs ~$0.04/image, and was immediately accessible with the existing key.",
    ),
    (
        "DynamoDB vs RDS",
        "DynamoDB PAY_PER_REQUEST",
        "Aurora Serverless PostgreSQL",
        "Campaign data is accessed by primary key 99% of the time. "
        "DynamoDB delivers microsecond reads with no query language required. "
        "PAY_PER_REQUEST means zero cost at idle — Aurora's minimum is ~$50/month even at zero traffic.",
    ),
    (
        "Polling vs WebSockets",
        "Client polls GET /campaigns/{id} every 3s",
        "API Gateway WebSocket + Lambda push",
        "Generation takes 20–30 seconds. Polling at 3s costs ~10 API calls per campaign — negligible. "
        "WebSockets require connection management, reconnection logic, a DynamoDB connection table, "
        "and significantly more infrastructure. The UX difference is imperceptible for a 30-second job.",
    ),
    (
        "HTTP API v2 vs REST API v1",
        "HTTP API v2 with Cognito JWT native authoriser",
        "REST API v1 with API key authentication",
        "HTTP v2 supports Cognito JWT authorisers natively — no Lambda authoriser needed. "
        "It also costs 3.5× less ($1.00/M vs $3.50/M). Using Cognito means the reviewer's "
        "identity is extracted from the JWT server-side; the client cannot spoof it.",
    ),
]

for title, chose, alt, reason in decisions:
    h3(title)
    add_table(
        ["", ""],
        [
            ["Chose", chose],
            ["Alternative", alt],
            ["Reason", reason],
        ]
    )

# ═══════════════════════════════════════════════════════════════════════════════
# 8. HONEST TRADE-OFFS
# ═══════════════════════════════════════════════════════════════════════════════
h1("8. Honest Trade-offs")

tradeoffs = [
    (
        "A. Imagen 4 has no reference photo seeding",
        "Nova Canvas supports IMAGE_VARIATION mode — using a product photo as a conditioning "
        "seed for brand-anchored generation. Imagen 4 is text-to-image only. "
        "Prompts describe the product but there is no pixel-level brand anchoring.",
        "Switch to Nova Canvas once Bedrock model access is activated."
    ),
    (
        "B. Japanese and CJK text overlay does not render",
        "The bundled RobotoSlab-Bold.ttf does not include CJK character sets. "
        "Japanese, Korean, and Chinese headlines cannot be composited as text overlays.",
        "See The Roadmap section — CJK font support is the first planned improvement."
    ),
    (
        "C. Compliance is a filter, not legal clearance",
        "GPT-4o mini applies a consistent rubric but is not a licensed legal reviewer. "
        "Pass/warn/fail verdicts are a first-pass filter. Human approval remains mandatory.",
        "The pending_review workflow enforces human sign-off before any asset goes live."
    ),
]

for title, limitation, mitigation in tradeoffs:
    h3(title)
    body(limitation)
    body(f"Mitigation: {mitigation}", italic=True)

# ═══════════════════════════════════════════════════════════════════════════════
# 9. THE ROADMAP
# ═══════════════════════════════════════════════════════════════════════════════
h1("9. The Roadmap")
h2("Short Term")

h3("Japanese and CJK Language Support")
body(
    "The current font does not contain Japanese, Korean, or Chinese glyphs. "
    "Supporting these markets requires bundling CJK-compatible fonts and "
    "selecting them at runtime based on the target language code:"
)
code(
    "FONT_MAP = {\n"
    '    "ja": "NotoSansJP-Bold.ttf",    # Japanese\n'
    '    "ko": "NotoSansKR-Bold.ttf",    # Korean\n'
    '    "zh": "NotoSansSC-Bold.ttf",    # Simplified Chinese\n'
    '    "default": "RobotoSlab-Bold.ttf"\n'
    "}\n"
    "lang = data.get('language', 'en')[:2]\n"
    "font_file = FONT_MAP.get(lang, FONT_MAP['default'])"
)
body(
    "GPT-4o mini already produces high-quality Japanese copy — this is purely a font rendering fix. "
    "It unblocks campaigns targeting Japan (~$3.7T consumer market) with no API or model changes."
)

h3("Reference-Led Image Generation")
body(
    "Switch to Nova Canvas IMAGE_VARIATION mode once Bedrock model access is enabled. "
    "Product reference photos uploaded to S3 serve as conditioning seeds, "
    "anchoring outputs to the actual product rather than text description."
)

h2("Long Term")
add_table(
    ["Improvement", "What it enables"],
    [
        ["Video generation (Nova Reel / Google Veo)", "15s and 30s social video assets with same compliance gates"],
        ["Ad platform integration (Meta, TikTok)", "One-click publish from approved state via platform APIs"],
        ["Fine-tuning on approved campaigns", "Learning loop on ~1,000 approved briefs — closes FR12"],
        ["Horizontal micro-scaling", "Fan-out to one SQS message per product for true horizontal scale"],
    ]
)

# ═══════════════════════════════════════════════════════════════════════════════
# QUICK REFERENCE
# ═══════════════════════════════════════════════════════════════════════════════
h1("Quick Reference")
h2("Stack Summary")
add_table(
    ["Stack", "Resources", "Purpose"],
    [
        ["Secrets", "Secrets Manager (2 secrets)", "OpenAI + Google AI API keys"],
        ["Storage", "S3 (outputs) + DynamoDB", "Images + campaign state"],
        ["Pipeline", "SQS + GenerateCampaign Lambda", "Async generation, maxConcurrency=10"],
        ["Api", "HTTP API v2 + Cognito + 2 Lambdas", "Authenticated brief submission + polling"],
    ]
)

h2("Brief Schema")
code(
    '{\n'
    '  "product_name": "LuxeRoll Hair Brush",\n'
    '  "region":       "united states",\n'
    '  "audience":     "working women aged 24-40, mid-to-high income",\n'
    '  "message":      "Define your natural beauty, on your schedule",\n'
    '  "language":     "en"\n'
    '}'
)

h2("Cost Model")
add_table(
    ["Scale", "Est. Monthly Cost"],
    [
        ["100 campaigns/month", "~$13"],
        ["500 campaigns/month", "~$65"],
        ["2,000 campaigns/month", "~$260"],
    ]
)

doc.save(OUT)
print(f"Saved: {OUT}")
